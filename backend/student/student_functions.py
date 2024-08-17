import os
import uuid
from datetime import datetime

import docx
from flask_jwt_extended import jwt_required
from werkzeug.utils import secure_filename

from app import app, api, request, jsonify, db, contains_eager, desc
from backend.functions.small_info import checkFile, user_contract_folder
from backend.functions.utils import find_calendar_date, update_week
from backend.models.models import Students, AttendanceHistoryStudent, DeletedStudents, Users, RegisterDeletedStudents, \
    Contract_Students, BookPayments, StudentPayments, Teachers, Roles, Locations, StudentExcuses, StudentHistoryGroups, \
    Groups, Contract_Students_Data, StudentCharity, GroupReason


@app.route(f'{api}/student_history2/<int:user_id>')
@jwt_required()
def student_history(user_id):
    years = []
    student_get = Students.query.filter(Students.user_id == user_id).first()
    history_rates = AttendanceHistoryStudent.query.filter(
        AttendanceHistoryStudent.student_id == student_get.id).order_by(desc(AttendanceHistoryStudent.id)).all()

    history_rate_list = [
        {
            "group_id": rate.group.id,
            "group_name": rate.group.name.title(),
            "subject": rate.subject.name,
            "degree": rate.average_ball,
            "month": rate.month.date.strftime("%h"),
            "year": rate.year.date.strftime("%Y"),
            "teacher_id": rate.group.teacher_id,
            "teacher_name": Teachers.query.filter(Teachers.id == rate.group.teacher_id).first().user.name,
            "teacher_surname": Teachers.query.filter(Teachers.id == rate.group.teacher_id).first().user.surname,
        } for rate in history_rates
    ]
    years = [rate.year.date.strftime("%Y") for rate in history_rates]
    student_groups = StudentHistoryGroups.query.filter(StudentHistoryGroups.student_id == student_get.id).order_by(
        desc(StudentHistoryGroups.id)).all()
    history_group_list = [
        {
            "group_id": gr.group.id,
            "group_name": gr.group.name.title(),
            "reason": gr.reason,
            "joined_day": gr.joined_day.strftime("%Y-%m-%d"),
            "left_day": gr.left_day.strftime("%Y-%m-%d") if gr.left_day else "",
            "teacher_id": gr.group.teacher_id,
            "teacher_name": Teachers.query.filter(Teachers.id == gr.group.teacher_id).first().user.name.title(),
            "teacher_surname": Teachers.query.filter(Teachers.id == gr.group.teacher_id).first().user.surname.title(),
        } for gr in student_groups
    ]
    years = list(dict.fromkeys(years))
    return jsonify({
        "data": {
            "history_rate": history_rate_list,
            "years": years,
            "history_groups": history_group_list
        }
    })


@app.route(f'{api}/add_blacklist/<int:user_id>', methods=["GET", "POST"])
@jwt_required()
def add_blacklist2(user_id):
    user = Users.query.filter(Users.id == user_id).first()
    calendar_year, calendar_month, calendar_day = find_calendar_date()
    student = Students.query.filter(Students.user_id == user_id).first()
    if student.debtor != 4:
        student.debtor = 4
        db.session.commit()
        return jsonify({
            "success": True,
            "msg": "Student qora ro'yxatga qo'shildi"
        })

    else:
        student_excuse = StudentExcuses.query.filter(StudentExcuses.student_id == student.id,
                                                     StudentExcuses.to_date > calendar_day.date).order_by(
            desc(StudentExcuses.id)).first()

        if not student_excuse:
            if user.balance >= 0:
                Students.query.filter_by(id=student.id).update({"debtor": 0})
            elif user.balance < 0:
                Students.query.filter_by(id=student.id).update({"debtor": 1})
            elif user.balance >= -student.combined_debt:
                Students.query.filter_by(id=student.id).update({"debtor": 2})
        else:
            Students.query.filter_by(id=student.id).update({"debtor": 3})
        db.session.commit()
        return jsonify({
            "success": True,
            "msg": "Student qora ro'yxatdan olib tashlandi"
        })


@app.route(f'{api}/delete_newStudent/<int:user_id>', methods=["GET", "POST"])
@jwt_required()
def delete_newStudent(user_id):
    calendar_year, calendar_month, calendar_day = find_calendar_date()
    student = Students.query.filter(Students.user_id == user_id).first()
    reason = request.get_json()['reason']
    del_new_student = RegisterDeletedStudents(student_id=student.id, reason=reason, calendar_day=calendar_day)
    db.session.add(del_new_student)
    db.session.commit()
    return jsonify({
        "success": True,
        "msg": "Student ro'yxatdan o'chirildi"
    })


@app.route(f'{api}/get_back_student/<int:user_id>')
@jwt_required()
def get_back_student(user_id):
    student = Students.query.filter(Students.user_id == user_id).first()
    del_new_student = RegisterDeletedStudents.query.filter(RegisterDeletedStudents.student_id == student.id).first()
    db.session.delete(del_new_student)
    db.session.commit()
    return jsonify({
        "success": True,
        "msg": "Student ro'yxatga qaytarildi"
    })


@app.route(f'{api}/studyingStudents/<int:id>', methods=['POST', 'GET'])
@jwt_required()
def studyingStudents(id):
    user_list = Users.query.join(Students).filter(Students.group != None, Users.location_id == id).order_by(
        Users.id).all()
    user_id = []
    for user in user_list:
        user_id.append(user.id)
    user_id = list(dict.fromkeys(user_id))

    students_list = Students.query.filter(Students.user_id.in_([user_id for user_id in user_id])).join(
        Students.group).filter(Groups.status == True).order_by(Students.user_id).all()
    role = Roles.query.filter(Roles.type_role == "student").first()

    list_students = [
        {
            "id": st.user.id,
            "name": st.user.name.title(),
            "surname": st.user.surname.title(),
            "username": st.user.username,
            "language": st.user.language.name,
            "age": st.user.age,
            "reg_date": st.user.day.date.strftime("%Y-%m-%d"),
            "comment": st.user.comment,
            'money': st.user.balance,
            "role": role.role,
            "subjects": [sub.name for sub in st.subject],
            "photo_profile": st.user.photo_profile,
            "moneyType": ["green", "yellow", "red", "navy", "black"][st.debtor] if st.debtor else 0
        } for st in students_list
    ]
    return jsonify({
        "studyingStudents": list_students
    })


@app.route(f'{api}/deletedStudents/<int:id>', methods=['POST'])
@jwt_required()
def deletedStudents(id):
    reason = request.get_json()['type']
    user_list = db.session.query(Students).join(Students.user).options(contains_eager(Students.user)).filter(
        Students.deleted_from_group != None, Users.location_id == id).order_by('id').all()
    user_id = []
    for user in user_list:
        user_id.append(user.id)
    user_id = list(dict.fromkeys(user_id))

    if reason == "Hammasi":
        students_list = DeletedStudents.query.filter(
            DeletedStudents.student_id.in_([user_id for user_id in user_id])).order_by(
            desc(DeletedStudents.calendar_day)).all()
    else:
        group_reason = GroupReason.query.filter(GroupReason.id == reason).first()
        students_list = DeletedStudents.query.filter(DeletedStudents.student_id.in_([user_id for user_id in user_id]),
                                                     DeletedStudents.reason_id == group_reason.id).order_by(
            desc(DeletedStudents.calendar_day)).all()

    role = Roles.query.filter(Roles.type_role == "student").first()

    list_students = [
        {
            "id": st.student.user.id,
            "name": st.student.user.name.title(),
            "surname": st.student.user.surname.title(),
            "username": st.student.user.username,
            "language": st.student.user.language.name,
            "age": st.student.user.age,
            "reg_date": st.student.user.day.date.strftime("%Y-%m-%d"),
            "deleted_date": st.day.date.strftime("%Y-%m-%d"),
            "day": st.calendar_day,
            "teacher": st.teacher_id,
            "comment": st.student.user.comment,
            'money': st.student.user.balance,
            "role": role.role,
            "photo_profile": st.student.user.photo_profile,
            "moneyType": ["green", "yellow", "red", "navy", "black"][st.student.debtor] if st.student.debtor else 0,
            "phone": st.student.user.phone[0].phone,
            "reason": st.reason,
            "group": st.group.id
        }
        for st in students_list
    ]
    day_dict = {gr['id']: gr for gr in list_students}
    day_list = list(day_dict.values())

    return jsonify({
        "data": day_list
    })


@app.route(f"{api}/newStudents/<int:location_id>", methods=["GET"])
@jwt_required()
def newStudents(location_id):
    update_week(location_id)
    students = Students.query.join(Users).filter(Users.location_id == location_id, Users.student != None,
                                                 Students.subject != None,
                                                 Students.deleted_from_register == None).order_by(
        desc(Students.id)).all()
    list_students = [
        st.convert_json() for st in students
    ]
    return jsonify({
        "newStudents": list_students
    })


@app.route(f'{api}/get_filtered_students_list/<int:location_id>', methods=["GET"])
@jwt_required()
def get_filtered_students_list(location_id):
    students = Students.query.join(Users).filter(Users.location_id == location_id, Users.student != None,
                                                 Students.subject != None,
                                                 Students.deleted_from_register == None).order_by(
        desc(Students.id)).all()
    subjects_with_students = {}

    for student in students:
        for subject in student.subject:
            if subject.id not in subjects_with_students:
                subjects_with_students[subject.id] = {
                    "id": subject.id,
                    "name": subject.name,
                    "students": []
                }
            subjects_with_students[subject.id]["students"].append(student.convert_json())
    return jsonify(list(subjects_with_students.values()))


@app.route(f'{api}/new_del_students/<location_id>')
@jwt_required()
def newStudents_deleted(location_id):
    role = Roles.query.filter(Roles.type_role == "student").first()
    students = db.session.query(Users).join(Users.student).options(contains_eager(Users.student)).filter(
        Users.location_id == location_id, Users.student != None, Students.deleted_from_register != None,
    ).join(Users.day).options(contains_eager(Users.month)).order_by(
        desc(Users.id)).all()
    students = Students.query.join(Users).filter(Users.location_id == location_id, Users.student != None,
                                                 Students.subject != None,
                                                 Students.deleted_from_register != None).order_by(
        desc(Students.id)).all()
    subjects_with_students = {}
    for student in students:
        for subject in student.subject:
            if subject.id not in subjects_with_students:
                subjects_with_students[subject.id] = {
                    "id": subject.id,
                    "name": subject.name,
                    "students": []
                }
            subjects_with_students[subject.id]["students"].append(student.convert_json())
    return jsonify(list(subjects_with_students.values()))


@app.route(f'{api}/create_contract/<int:user_id>', methods=["POST"])
@jwt_required()
def create_contract(user_id):
    calendar_year, calendar_month, calendar_day = find_calendar_date()
    name = request.get_json()['name']
    surname = request.get_json()['surname']
    passportSeries = request.get_json()['passportSeries']
    fatherName = request.get_json()['fatherName']
    givenPlace = request.get_json()['givenPlace']
    givenTime = request.get_json()['givenTime']
    place = request.get_json()['place']
    Students.query.filter(Students.user_id == user_id).update({
        "representative_name": name,
        "representative_surname": surname
    })
    db.session.commit()

    ot = request.get_json()['date']['ot']
    do = request.get_json()['date']['do']
    ot = datetime.strptime(ot, "%Y-%m-%d")
    do = datetime.strptime(do, "%Y-%m-%d")
    ot_month = datetime.strftime(ot, "%m")
    do_month = datetime.strftime(do, "%m")
    do_year = datetime.strftime(do, "%Y")
    if int(do_year) > int(calendar_year.date.strftime("%Y")):
        do_month = int(do_month) + 12
    month = int(do_month) - int(ot_month) + 1
    ot = datetime.strftime(ot, "%Y-%m-%d")
    do = datetime.strftime(do, "%Y-%m-%d")

    user = Users.query.filter(Users.id == user_id).first()

    student = Students.query.filter(Students.user_id == user_id).first()
    location = Locations.query.filter(Locations.id == user.location_id).first()
    contract = Contract_Students.query.filter(Contract_Students.student_id == student.id).first()

    if student.contract_word_url:
        if os.path.exists(student.contract_word_url):
            os.remove(student.contract_word_url)
    Students.query.filter(Students.id == student.id).update({
        "contract_word_url": ""
    })
    db.session.commit()
    student_charity = StudentCharity.query.filter(StudentCharity.student_id == student.id).all()
    all_charity = 0
    for char in student_charity:
        all_charity += char.discount
    print('all charity', all_charity)
    contract_data = Contract_Students_Data.query.filter(Contract_Students_Data.location_id == location.id,
                                                        Contract_Students_Data.year == calendar_year.date).first()
    if not contract:
        contract = Contract_Students(student_id=student.id, created_date=ot,
                                     expire_date=do, father_name=fatherName, given_place=givenPlace,
                                     place=place, passport_series=passportSeries, given_time=givenTime)
        db.session.add(contract)
        db.session.commit()

        if not contract_data:
            new = Contract_Students_Data(year=calendar_year.date, number=1, location_id=location.id)
            db.session.add(new)
            db.session.commit()
        else:
            contract_data.number += 1
            db.session.commit()
    else:
        Contract_Students.query.filter(Contract_Students.student_id == student.id,
                                       ).update({
            "created_date": ot,
            "expire_date": do,
            "father_name": fatherName,
            "given_place": givenPlace,
            "place": place,
            "passport_series": passportSeries,
            "given_time": givenTime

        })
        db.session.commit()
    if user.age >= 18:
        name = user.name
        surname = user.surname
        father_name = user.father_name
    else:
        name = student.representative_name
        surname = student.representative_surname
        father_name = contract.father_name
    doc = docx.Document('frontend/build/static/contract_folder/contract.docx')
    id = uuid.uuid1()
    text = location.address.split(" ")
    text_item = ""
    text_item2 = ""
    if len(text) > 3:
        for item in text[0:4]:
            text_item += f" {item}"
        for item in text[4:]:
            text_item2 += f" {item}"
    else:
        for item in text:
            text_item += f" {item}"
    user_id = id.hex[0:15]
    campus_name = location.name + " " + location.location_type if location.location_type == "Shahri" else location.district + " " + location.location_type
    # number = f'{calendar_year.date.strftime("%Y")}/{contract_data.location.code}-{contract_data.number}'
    number = f'{calendar_year.date.strftime("%Y")}/{location.code}/{contract_data.number}'
    doc.paragraphs[0].runs[0].text = f"SHARTNOMA N{number}"
    doc.paragraphs[
        3].text = f"              {campus_name}			                                             {contract.created_date.strftime('%d-%m-%Y')}"
    print(doc.paragraphs[4].text)
    if location.id > 3:

        doc.paragraphs[
            5].text = "Oʻzbekiston Respublikasi Prezidentining 15.09.2017-yildagi PQ-3276 sonli “Nodavlat taʼlim xizmatlari koʻrsatish faoliyatini yanada rivojlantirish chora-tadbirlari toʻgʻrisida”gi qaroriga."
        doc.paragraphs[
            6].text = f"Hamda taʼlim muassasasi Ustaviga asosan faoliyat yurituvchi “{location.campus_name}” nodavlat taʼlim muassasasi (kelgusida “Nodavlat taʼlim muassasasi” deb yuritiluvchi)  nomidan direktor {location.director_fio} bir tomondan va {surname.title()} {name.title()} {father_name[0].title()}{father_name[1:].lower()}"
    else:
        doc.paragraphs[
            6].text = f"№ MTT 0428 Litsenziyaga asosan hamda taʼlim muassasasi Ustaviga asosan faoliyat yurituvchi “{location.campus_name}” nodavlat taʼlim muassasasi (kelgusida “Nodavlat taʼlim muassasasi” deb yuritiluvchi)  nomidan direktor {location.director_fio} bir tomondan va {surname.title()} {name.title()} {father_name[0].title()}{father_name[1:].lower()}"
    doc.paragraphs[
        9].text = f"1.1 Mazkur shartnomaga asosan oʻquvchining ota-onasi (yoki qonuniy vakili) nodavlat taʼlim muassasasiga maktabdan tashqari taʼlim olish maqsadida oʻzining voyaga yetmagan farzandi  {user.name.title()} {user.surname.title()} {user.father_name[0].title()}{user.father_name[1:].lower()} ni"

    doc.paragraphs[
        15].text = f"2.1. Oʻquvchining nodavlat taʼlim muassasasida taʼlim olishi uchun bir oylik toʻlov summasi {abs(student.combined_debt)} va {contract.expire_date.strftime('%d-%m-%Y')} muddatgacha {abs(((student.combined_debt) - all_charity) * month)}  soʻmni tashkil etadi."
    doc.paragraphs[
        69].text = f"7.1.Mazkur shartnoma tomonlar oʻrtasida imzolangan kundan boshlab yuridik kuchga ega hisoblanadi va {contract.expire_date.strftime('%d-%m-%Y')} muddatga qadar amal qiladi"
    info = [
        {
            "left_info": f"{location.campus_name} NTM",
            "right_info": f"F.I.O : {surname.title()} {name.title()} {father_name[0].title()}{father_name[1:].lower()}"
        },
        {
            "left_info": location.address,
            "right_info": f"Pasport maʼlumoti: Seriya {contract.passport_series}"
        },
        {
            "left_info": f"R/S: {location.bank_sheet}  INN: {location.inn}",
            "right_info": f"Berilgan vaqti: {contract.given_time}"
        },
        {
            "left_info": f"Bank: {location.bank}",
            "right_info": f"Manzili: {contract.place}"
        },
        {
            "left_info": f"MFO: {location.mfo}",
            "right_info": ""
        },
        {
            "left_info": f"Tel: {location.number_location}",
            "right_info": ""
        },
        {
            "left_info": f"Direktor: __________{location.director_fio}",
            "right_info": ""
        },
        {
            "left_info": "M.P",
            "right_info": "Imzo____________"
        },
    ]
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nodavlat taʼlim muassasasi'
    hdr_cells[1].text = 'O`quvchining qonuniy vavakili(ota-onasi)'
    for item in info:
        row_cells = table.add_row().cells
        row_cells[0].text = item['left_info']
        row_cells[1].text = item['right_info']

    doc.save(
        f"frontend/build/static/contract_folder/{user_id} {student.user.name.title()} {student.user.surname.title()}doc.docx")
    new_doc = f"static/contract_folder/{user_id} {student.user.name.title()} {student.user.surname.title()}doc.docx"
    Students.query.filter(Students.id == student.id).update({
        "contract_word_url": f"frontend/build/static/contract_folder/{user_id} {student.user.name.title()} {student.user.surname.title()}doc.docx"
    })
    db.session.commit()
    return jsonify({
        "success": True,
        "msg": "Shartnoma yaratildi",
        "file": new_doc
    })


@app.route(f'{api}/upload_pdf_contract/<int:user_id>', methods=["POST"])
@jwt_required()
def upload_pdf_contract(user_id):
    student = Students.query.filter(Students.user_id == user_id).first()
    file = request.files['file']
    app.config['UPLOAD_FOLDER'] = user_contract_folder()

    url = ""
    if file and checkFile(file.filename):
        file.filename = f"{student.id}/{student.user.name}/{student.user.surname}.pdf"
        file_name = secure_filename(file.filename)
        file.save(os.path.join(app.config['UPLOAD_FOLDER'], file_name))
        url = "static" + "/" + "contract_pdf" + "/" + file_name
    Students.query.filter(Students.user_id == user_id).update({
        "contract_pdf_url": url
    })
    db.session.commit()
    return jsonify({
        "success": True,
        "msg": "Fayl yuklandi",
        "url": url
    })


@app.route(f'{api}/change_location/<int:user_id>/<int:location_id>')
@jwt_required()
def change_location(user_id, location_id):
    location = Locations.query.filter(Locations.id == location_id).first()
    user = Users.query.filter(Users.id == user_id).first()
    if user.location_id != location_id:
        Users.query.filter(Users.id == user_id).update({
            "location_id": location_id
        })
        db.session.commit()
        return jsonify({
            "msg": f"O'quvchi {location.name} flialiga qo'shildi",
            "success": True
        })
    else:
        return jsonify({
            "msg": f"O'quvchi allaqachon {location.name} fliada",
            "success": True
        })


@app.route(f"{api}/student_attendance_info/<user_id>")
@jwt_required()
def student_attendance_info(user_id):
    student = Students.query.filter(Students.user_id == user_id).first()
    attendance_histories = AttendanceHistoryStudent.query.filter(
        AttendanceHistoryStudent.student_id == student.id).order_by(AttendanceHistoryStudent.id).all()
    student_payments = StudentPayments.query.filter(StudentPayments.student_id == student.id,
                                                    StudentPayments.payment == True).order_by(
        StudentPayments.id).all()
    history_list = []
    book_payments = BookPayments.query.filter(BookPayments.student_id == student.id).order_by(
        BookPayments.id).all()

    book_payment_list = [
        {
            "id": bk_payment.id,
            "payment": bk_payment.payment_sum,
            "date": bk_payment.day.date.strftime("%Y-%m-%d")
        } for bk_payment in book_payments
    ]
    history_list = [
        {
            "group_name": att.group.subject.name if att.group else "Ma'lumot yo'q",
            "total_debt": att.total_debt,
            "payment": att.payment,
            "remaining_debt": att.remaining_debt,
            "discount": att.total_discount,
            "present": att.present_days + att.scored_days,
            "absent": att.absent_days,
            "days": att.present_days + att.absent_days,
            "month": att.month.date.strftime("%Y-%m")
        } for att in attendance_histories
    ]
    payment_list = [
        {
            "id": payment.id,
            "payment": payment.payment_sum,
            "date": payment.day.date.strftime("%Y-%m-%d"),
            "type_payment": payment.payment_type.name
        } for payment in student_payments
    ]

    student_payments = StudentPayments.query.filter(StudentPayments.student_id == student.id,
                                                    StudentPayments.payment == False).order_by(
        StudentPayments.id).all()
    discount_list = [
        {
            "id": payment.id,
            "payment": payment.payment_sum,
            "date": payment.day.date.strftime("%Y-%m-%d"),

        } for payment in student_payments
    ]
    return jsonify({
        "data": {
            "id": student.user.id,
            "name": student.user.name.title(),
            "surname": student.user.surname.title(),
            "debts": history_list,
            "payments": payment_list,
            "discounts": discount_list,
            "bookPayments": book_payment_list
        }
    })

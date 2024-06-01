from app import app, db, jsonify, contains_eager, request, desc, send_file
from werkzeug.utils import secure_filename

from backend.models.models import Overhead, AccountingPeriod, PaymentTypes, DeletedOverhead, \
    CalendarMonth, CalendarDay, Category, ConnectedCategory, Capital, CapitalTerm
from flask_jwt_extended import jwt_required
from backend.functions.filters import old_current_dates
from backend.functions.small_info import room_images, checkFile
from backend.functions.utils import get_json_field, find_calendar_date, api, iterate_models, get_form_field
from .utils import update_capital
from datetime import datetime
import os
from docx import Document
from pprint import pprint
import json
from backend.models.models import func


@app.route(f'{api}/add_capital_category', methods=['POST', "PUT", "DELETE"])
@jwt_required()
def add_capital_category():
    if request.method == "DELETE":
        category_id = get_json_field('category_id')
        connected_category = ConnectedCategory.query.filter(
            ConnectedCategory.addition_category_id == category_id).first()
        db.session.delete(connected_category)
        db.session.commit()
        Category.query.filter(Category.id == category_id).delete()
        db.session.commit()
        return jsonify({
            "msg": "O'chirildi",
            "success": True
        })
    info = request.form.get("info")
    json_file = json.loads(info)
    if 'category_id' in json_file:
        category_id = json_file['category_id']
    else:
        category_id = None
    name = json_file['name']
    number_category = json_file['number_category']
    img = request.files.get('img')
    url = ""
    if img and checkFile(img.filename):
        app.config['UPLOAD_FOLDER'] = room_images()
        photo_filename = secure_filename(img.filename)
        img.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))
        url = "static" + "/" + "room" + "/" + photo_filename
    if request.method == "POST":

        if not category_id:
            new = Category(name=name, number=1, number_category=number_category, img=url)
            db.session.add(new)
            db.session.commit()

        else:
            category = Category.query.filter(Category.id == category_id).first()
            new = Category(name=name, number=category.number + 1, number_category=number_category, img=url)
            db.session.add(new)
            db.session.commit()
            new_connected = ConnectedCategory(addition_category_id=new.id, main_category_id=category_id)
            db.session.add(new_connected)
            db.session.commit()
        return jsonify({
            'status': True,
            "category": new.convert_json(),
            "msg": f"{new.name} turi yaratildi"
        })
    else:
        category = Category.query.filter(Category.id == category_id).first()
        category.name = name
        category.number_category = number_category
        if url:
            category.img = url
        db.session.commit()
        return jsonify({
            'status': True,
            "category": category.convert_json(),
            "msg": f"{category.name} o'zgartirildi"
        })


@app.route(f'{api}/get_capital_category/<int:category_id>/<int:location_id>', methods=['GET', "POST"])
@jwt_required()
def get_capital_category(category_id, location_id):
    category = Category.query.filter(Category.id == category_id).first()
    update_capital(location_id)

    if request.method == "POST":
        info = request.form.get("info")
        json_file = json.loads(info)
        name = json_file['name']
        number_category = json_file['number_category']
        img = request.files.get('img')
        if img and checkFile(img.filename):
            app.config['UPLOAD_FOLDER'] = room_images()
            photo_filename = secure_filename(img.filename)
            img.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))
            url = "static" + "/" + "room" + "/" + photo_filename
            category.img = url
            db.session.commit()
        category.name = name
        category.number_category = number_category
        db.session.commit()
        return jsonify({
            "msg": f"{category.name} ma'lumotlari o'zgartirildi",
            'category': category.convert_json(location_id)
        })
    return jsonify({
        'category': category.convert_json(location_id),
        "capital_tools": old_current_dates(),
    })


@app.route(f'{api}/deleted_capitals/<int:category_id>/<int:location_id>')
@jwt_required()
def deleted_capitals(category_id, location_id):
    capitals = Capital.query.filter(Capital.category_id == category_id, Capital.location_id == location_id).order_by(
        Capital.id).all()
    return jsonify({
        "capitals": iterate_models(capitals)
    })


@app.route(f'{api}/get_capital_categories/<int:location_id>', methods=['GET'])
@jwt_required()
def get_capital_categories(location_id):
    categories = Category.query.filter(Category.number == 1).order_by(Category.id).all()
    list = []
    total_down_cost = 0
    for category in categories:
        list.append(category.convert_json())
        all_capex_down = \
            db.session.query(func.sum(Capital.total_down_cost).filter(Capital.category_id == category.id,
                                                                      Capital.location_id == location_id,
                                                                      Capital.deleted != True)).first()[0]
        total_down_cost += all_capex_down if all_capex_down else  0

    return jsonify({
        'categories': list,
        "total_down_cost": total_down_cost
    })


@app.route(f'{api}/add_capital', defaults={"location_id": None}, methods=['POST', "PUT", "DELETE"])
@app.route(f'{api}/add_capital/<location_id>', methods=['POST', "PUT", "DELETE"])
def add_capital(location_id):
    accounting_period = db.session.query(AccountingPeriod).join(AccountingPeriod.month).options(
        contains_eager(AccountingPeriod.month)).order_by(desc(CalendarMonth.id)).first()
    if request.method == "POST":
        current_year = datetime.now().year
        old_year = datetime.now().year - 1
        month = str(datetime.now().month)
        info = request.form.get("info")
        json_file = json.loads(info)
        month_get = json_file.get('month')
        day = json_file.get('day')
        if month_get == "12" and month == "01":
            current_year = old_year
        if not month_get:
            month_get = month

        date_day = str(current_year) + "-" + str(month_get) + "-" + str(day)
        date_month = str(current_year) + "-" + str(month_get)
        date_year = datetime.strptime(str(current_year), "%Y")
        date_day = datetime.strptime(date_day, "%Y-%m-%d")
        date_month = datetime.strptime(date_month, "%Y-%m")
        calendar_year, calendar_month, calendar_day = find_calendar_date(date_day=date_day, date_month=date_month,
                                                                         date_year=date_year)
        img = request.files.get('img')
        url = ""
        if img and checkFile(img.filename):
            app.config['UPLOAD_FOLDER'] = room_images()
            photo_filename = secure_filename(img.filename)
            img.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))
            url = "static" + "/" + "room" + "/" + photo_filename

            db.session.commit()

        capital = {
            "name": json_file.get('capital_name'),
            "number": json_file.get('number'),
            "price": json_file.get('price'),
            "term": json_file.get('term'),
            "category_id": json_file.get('category_id'),
            "calendar_year": calendar_year.id,
            "location_id": location_id,
            "calendar_day": calendar_day.id,
            "account_period_id": accounting_period.id,
            "payment_type_id": json_file.get('payment_type_id'),
            "calendar_month": calendar_month.id,
            "img": url

        }

        capital_add = Capital(**capital)
        capital_add.add()
        return jsonify({
            "msg": f"{json_file.get('name')} ro'yxatga qo'shildi",
            "success": True,
            "capital": capital_add.convert_json()
        })
    elif request.method == "PUT":
        current_year = datetime.now().year
        old_year = datetime.now().year - 1
        month = str(datetime.now().month)
        info = request.form.get("info")
        json_file = json.loads(info)
        month_get = json_file.get('month')
        day = json_file.get('day')
        if month_get == "12" and month == "01":
            current_year = old_year
        if not month_get:
            month_get = month

        date_day = str(current_year) + "-" + str(month_get) + "-" + str(day)
        date_month = str(current_year) + "-" + str(month_get)
        date_year = datetime.strptime(str(current_year), "%Y")
        date_day = datetime.strptime(date_day, "%Y-%m-%d")
        date_month = datetime.strptime(date_month, "%Y-%m")
        calendar_year, calendar_month, calendar_day = find_calendar_date(date_day=date_day, date_month=date_month,
                                                                         date_year=date_year)
        capital_id = json_file.get('capital_id')
        capital = Capital.query.filter(Capital.id == capital_id).first()
        capital.name = json_file.get('capital_name')
        capital.number = json_file.get('number')
        capital.price = json_file.get('price')
        capital.term = json_file.get('term')
        capital.calendar_day = calendar_day.id
        capital.calendar_month = calendar_month.id
        capital.calendar_day = calendar_day.id
        capital.payment_type_id = json_file.get('payment_type_id')
        img = request.files.get('img')

        if img and checkFile(img.filename):
            print(True)
            app.config['UPLOAD_FOLDER'] = room_images()
            photo_filename = secure_filename(img.filename)
            img.save(os.path.join(app.config['UPLOAD_FOLDER'], photo_filename))
            url = "static" + "/" + "room" + "/" + photo_filename
            capital.img = url
        db.session.commit()
        return jsonify({
            "msg": f"{capital.name} ma'lumotlari o'zgartirildi",
            "success": True,
            "capital": capital.convert_json()
        })
    else:
        capital_id = get_json_field('capital_id')
        capital = Capital.query.filter(Capital.id == capital_id).first()
        capital.deleted = True
        db.session.commit()
        return jsonify({
            "msg": f"{capital.name} ro'yxatdan o'chirildi",
            "success": True
        })


@app.route(f'{api}/capital_info/<int:capital_id>')
@jwt_required()
def capital_info(capital_id):
    capital = Capital.query.filter(Capital.id == capital_id).first()
    capital_terms = CapitalTerm.query.filter(CapitalTerm.capital_id == capital_id).order_by(CapitalTerm.id).all()
    return jsonify({
        "terms": iterate_models(capital_terms),
        "capital": capital.convert_json(),
        "capital_tools": old_current_dates()
    })


@app.route(f'{api}/get_capital_numbers', methods=['POST'])
@jwt_required()
def get_capital_numbers():
    category_id = get_json_field('category_id')
    category = Category.query.filter(Category.id == category_id).first()
    capitals = Capital.query.filter(Capital.category_id == category.id, Capital.deleted != True).order_by(
        Capital.id).all()

    document = Document()

    document.add_heading(f'{category.name} ga kiruvchi buyumlar:', 0)
    print(capitals)
    table = document.add_table(rows=len(capitals) - 1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Number'
    num = 0
    for capital in capitals:
        row_cells = table.add_row().cells

        row_cells[0].text = capital.name
        row_cells[1].text = capital.number
    document.add_page_break()

    document.save(f'frontend/build/static/contract_folder/{category.name}.docx')
    return send_file(f"frontend/build/static/contract_folder/{category.name}.docx", mimetype=None, as_attachment=False,
                     download_name=category.name, conditional=True, etag=True,
                     last_modified=None, max_age=None)


@app.route(f'{api}/add_overhead/<int:location_id>', methods=['POST'])
@jwt_required()
def add_overhead(location_id):
    """
    add data to Overhead table
    :param location_id: Location table primary key
    :return:
    """
    calendar_year, calendar_month, calendar_day = find_calendar_date()
    account_period = db.session.query(AccountingPeriod).join(AccountingPeriod.month).options(
        contains_eager(AccountingPeriod.month)).order_by(desc(CalendarMonth.id)).first()
    current_year = datetime.now().year
    old_year = datetime.now().year - 1
    month = str(datetime.now().month)
    month_get = get_json_field('month')
    day = get_json_field('day')
    if month_get == "12" and month == "01":
        current_year = old_year
    if not month_get:
        month_get = month

    date_day = str(current_year) + "-" + str(month_get) + "-" + str(day)
    date_month = str(current_year) + "-" + str(month_get)
    date_year = str(current_year)
    date_day = datetime.strptime(date_day, "%Y-%m-%d")
    date_month = datetime.strptime(date_month, "%Y-%m")
    calendar_month = CalendarMonth.query.filter(CalendarMonth.date == date_month).first()
    calendar_day = CalendarDay.query.filter(CalendarDay.date == date_day).first()
    type_of_data = get_json_field('typePayment')
    sum = int(get_json_field('price'))
    name_item = get_json_field('typeItem')
    payment_type = PaymentTypes.query.filter(PaymentTypes.id == type_of_data).first()

    add = Overhead(item_sum=sum, item_name=name_item, payment_type_id=payment_type.id, location_id=location_id,
                   calendar_day=calendar_day.id, calendar_month=calendar_month.id, calendar_year=calendar_year.id,
                   account_period_id=account_period.id)
    db.session.add(add)
    db.session.commit()
    return jsonify({
        "success": True,
        "msg": "Qo'shimcha xarajat qo'shildi"
    })


@app.route(f'{api}/delete_overhead/<overhead_id>', methods=["POST"])
@jwt_required()
def delete_overhead(overhead_id):
    """
    delete data from Overhead table
    :param overhead_id: Overhead table primary key
    :return:
    """
    reason = get_json_field('otherReason')
    calendar_year, calendar_month, calendar_day = find_calendar_date()

    overhead = Overhead.query.filter(Overhead.id == overhead_id).first()

    deleted_overhead = DeletedOverhead(item_sum=overhead.item_sum, item_name=overhead.item_name,
                                       payment_type_id=overhead.payment_type_id, location_id=overhead.location_id,
                                       calendar_day=overhead.calendar_day, calendar_month=overhead.calendar_month,
                                       calendar_year=overhead.calendar_year,
                                       account_period_id=overhead.account_period_id, deleted_date=calendar_day.date,
                                       reason=reason)
    db.session.add(deleted_overhead)
    db.session.commit()
    db.session.delete(overhead)
    db.session.commit()
    return jsonify({
        "success": True,
        "msg": "Xarajat o'chirildi"
    })


@app.route(f'{api}/change_overhead/<int:overhead>/<type_id>')
@jwt_required()
def change_overhead(overhead, type_id):
    """
    change payment_type_id in Overhead table
    :param overhead: Overhead table primary key
    :param type_id: Payment Type table primary key
    :return:
    """
    payment_type = PaymentTypes.query.filter(PaymentTypes.name == type_id).first()
    overhead_get = Overhead.query.filter(Overhead.id == overhead).first()

    Overhead.query.filter(Overhead.id == overhead_get.id).update({
        "payment_type_id": payment_type.id
    })
    db.session.commit()

    return jsonify({
        "success": True,
        "msg": "Xarajat summa turi o'zgartirildi"
    })

# @app.route(f'{api}/add_capital/<int:location_id>', methods=["GET", 'POST'])
# @jwt_required()
# def add_capital(location_id):
#     """
#     add tada to CapitalExpenditure table
#     :param location_id: Location table primary key
#     :return:
#     """
#     calendar_year, calendar_month, calendar_day = find_calendar_date()
#     if request.method == "POST":
#         account_period = db.session.query(AccountingPeriod).join(AccountingPeriod.month).options(
#             contains_eager(AccountingPeriod.month)).order_by(desc(CalendarMonth.id)).first()
#         current_year = datetime.now().year
#         month = get_json_field('month')
#         calendar_month = str(current_year) + "-" + str(month)
#         day = get_json_field('day')
#         day = str(current_year) + "-" + str(month) + "-" + str(day)
#         type_of_data = get_json_field('typePayment')
#         sum = int(get_json_field('price'))
#         name_item = get_json_field('typeItem')
#         payment_type = PaymentTypes.query.filter(PaymentTypes.id == type_of_data).first()
#         month = datetime.strptime(calendar_month, "%Y-%m")
#         day = datetime.strptime(day, "%Y-%m-%d")
#         calendar_month = CalendarMonth.query.filter(CalendarMonth.date == month).first()
#         calendar_day = CalendarDay.query.filter(CalendarDay.date == day).first()
#         add = CapitalExpenditure(item_sum=sum, item_name=name_item, payment_type_id=payment_type.id,
#                                  location_id=location_id,
#                                  calendar_day=calendar_day.id, calendar_month=calendar_month.id,
#                                  calendar_year=calendar_year.id,
#                                  account_period_id=account_period.id)
#         db.session.add(add)
#         db.session.commit()
#         return jsonify({
#             "success": True,
#             "msg": "Qo'shimcha xarajat qo'shildi"
#         })
#
#
# @app.route(f'{api}/delete_capital/<overhead_id>', methods=["POST"])
# @jwt_required()
# def delete_capital(overhead_id):
#     """
#     delete data from CapitalExpenditure table
#     :param overhead_id: CapitalExpenditure table primary key
#     :return:
#     """
#     reason = get_json_field('otherReason')
#     calendar_year, calendar_month, calendar_day = find_calendar_date()
#     capital = CapitalExpenditure.query.filter(CapitalExpenditure.id == overhead_id).first()
#
#     deleted_capital = DeletedCapitalExpenditure(item_sum=capital.item_sum, item_name=capital.item_name,
#                                                 payment_type_id=capital.payment_type_id,
#                                                 location_id=capital.location_id,
#                                                 calendar_day=capital.calendar_day,
#                                                 calendar_month=capital.calendar_month,
#                                                 calendar_year=capital.calendar_year,
#                                                 account_period_id=capital.account_period_id,
#                                                 deleted_date=calendar_day.date,
#                                                 reason=reason)
#     db.session.add(deleted_capital)
#     db.session.commit()
#     db.session.delete(capital)
#     db.session.commit()
#     return jsonify({
#         "success": True,
#         "msg": "Xarajat o'chirildi"
#     })
#
#
# @app.route(f'{api}/change_capital/<int:overhead>/<type_id>')
# @jwt_required()
# def change_capital(overhead, type_id):
#     """
#     change payment_type_id in Overhead table
#     :param overhead: CapitalExpenditure table primary key
#     :param type_id: Payment Type table primary key
#     :return:
#     """
#     payment_type = PaymentTypes.query.filter(PaymentTypes.name == type_id).first()
#     capital_get = CapitalExpenditure.query.filter(CapitalExpenditure.id == overhead).first()
#
#     CapitalExpenditure.query.filter(CapitalExpenditure.id == capital_get.id).update({
#         "payment_type_id": payment_type.id
#     })
#     db.session.commit()
#
#     return jsonify({
#         "success": True,
#         "msg": "Xarajat summa turi o'zgartirildi"
#     })

# @app.route(f'{api}/add_capital_category', methods=['POST'])
# def add_capital_category():
#     category = {
#         "name": get_json_field('name'),
#         "number": get_json_field('number')
#     }
#     capital_category = CapitalCategory(**category)
#     capital_category.add()
#     return jsonify({
#         "msg": f"{get_json_field('name')} kategoriyasi yaratildi",
#         "success": True
#     })
#
#
